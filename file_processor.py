# -*- coding: utf-8 -*-
"""
معالج الملفات لقراءة وتحليل جميع أنواع الملفات
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import PyPDF2
from docx import Document
import openpyxl
import json
import re
from datetime import datetime
import config

class FileProcessor:
    """معالج الملفات الرئيسي"""
    
    def __init__(self):
        """تهيئة معالج الملفات"""
        self.supported_formats = config.SUPPORTED_FORMATS
        self.max_file_size = config.MAX_FILE_SIZE
        
        # إنشاء المجلدات المطلوبة
        self.upload_folder = Path(config.UPLOAD_FOLDER)
        self.output_folder = Path(config.OUTPUT_FOLDER)
        self.upload_folder.mkdir(exist_ok=True)
        self.output_folder.mkdir(exist_ok=True)
    
    def get_file_type(self, file_path: str) -> str:
        """تحديد نوع الملف من امتداده"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            for file_type, extensions in self.supported_formats.items():
                if file_ext in extensions:
                    return file_type
            
            return "unknown"
        except Exception as e:
            logging.error(f"خطأ في تحديد نوع الملف: {e}")
            return "unknown"
    
    def validate_file(self, file_path: str, file_size: int) -> Tuple[bool, str]:
        """التحقق من صحة الملف"""
        try:
            # التحقق من حجم الملف
            if file_size > self.max_file_size:
                return False, f"حجم الملف كبير جداً. الحد الأقصى: {self.max_file_size // (1024*1024)} MB"
            
            # التحقق من وجود الملف
            if not os.path.exists(file_path):
                return False, "الملف غير موجود"
            
            # التحقق من نوع الملف
            file_type = self.get_file_type(file_path)
            if file_type == "unknown":
                return False, "نوع الملف غير مدعوم"
            
            return True, "الملف صالح"
            
        except Exception as e:
            logging.error(f"خطأ في التحقق من صحة الملف: {e}")
            return False, f"خطأ في التحقق من الملف: {str(e)}"
    
    def read_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """قراءة ملف PDF"""
        try:
            result = {
                "file_path": file_path,
                "file_type": "pdf",
                "content": "",
                "pages": 0,
                "metadata": {},
                "status": "success"
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # معلومات الملف
                result["pages"] = len(pdf_reader.pages)
                
                # استخراج النص من جميع الصفحات
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- الصفحة {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logging.warning(f"خطأ في قراءة الصفحة {page_num + 1}: {e}")
                        continue
                
                result["content"] = "\n\n".join(text_content)
                
                # استخراج البيانات الوصفية
                if pdf_reader.metadata:
                    result["metadata"] = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "producer": pdf_reader.metadata.get('/Producer', ''),
                        "creation_date": pdf_reader.metadata.get('/CreationDate', ''),
                        "modification_date": pdf_reader.metadata.get('/ModDate', '')
                    }
            
            return result
            
        except Exception as e:
            logging.error(f"خطأ في قراءة ملف PDF: {e}")
            return {
                "file_path": file_path,
                "file_type": "pdf",
                "content": "",
                "pages": 0,
                "metadata": {},
                "status": "error",
                "error": str(e)
            }
    
    def read_word_file(self, file_path: str) -> Dict[str, Any]:
        """قراءة ملف Word"""
        try:
            result = {
                "file_path": file_path,
                "file_type": "word",
                "content": "",
                "paragraphs": 0,
                "tables": 0,
                "metadata": {},
                "status": "success"
            }
            
            doc = Document(file_path)
            
            # استخراج النص من الفقرات
            paragraphs_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs_text.append(paragraph.text)
            
            result["content"] = "\n\n".join(paragraphs_text)
            result["paragraphs"] = len(paragraphs_text)
            
            # استخراج النص من الجداول
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    tables_text.append("\n".join(table_text))
            
            if tables_text:
                result["content"] += "\n\n--- الجداول ---\n" + "\n\n".join(tables_text)
                result["tables"] = len(doc.tables)
            
            # استخراج البيانات الوصفية
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or '',
                "author": core_props.author or '',
                "subject": core_props.subject or '',
                "keywords": core_props.keywords or '',
                "created": str(core_props.created) if core_props.created else '',
                "modified": str(core_props.modified) if core_props.modified else '',
                "revision": core_props.revision or 0
            }
            
            return result
            
        except Exception as e:
            logging.error(f"خطأ في قراءة ملف Word: {e}")
            return {
                "file_path": file_path,
                "file_type": "word",
                "content": "",
                "paragraphs": 0,
                "tables": 0,
                "metadata": {},
                "status": "error",
                "error": str(e)
            }
    
    def read_excel_file(self, file_path: str) -> Dict[str, Any]:
        """قراءة ملف Excel"""
        try:
            result = {
                "file_path": file_path,
                "file_type": "excel",
                "content": "",
                "sheets": [],
                "metadata": {},
                "status": "success"
            }
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # استخراج أسماء الأوراق
            sheet_names = workbook.sheetnames
            result["sheets"] = sheet_names
            
            # استخراج البيانات من كل ورقة
            sheets_data = []
            for sheet_name in sheet_names:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # استخراج البيانات من الصفوف
                for row in sheet.iter_rows(values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None and str(cell_value).strip():
                            row_data.append(str(cell_value).strip())
                    if row_data:
                        sheet_data.append(" | ".join(row_data))
                
                if sheet_data:
                    sheets_data.append(f"--- {sheet_name} ---\n" + "\n".join(sheet_data))
            
            result["content"] = "\n\n".join(sheets_data)
            
            # استخراج البيانات الوصفية
            result["metadata"] = {
                "sheets_count": len(sheet_names),
                "sheet_names": sheet_names,
                "properties": {
                    "title": workbook.properties.title or '',
                    "creator": workbook.properties.creator or '',
                    "subject": workbook.properties.subject or '',
                    "keywords": workbook.properties.keywords or '',
                    "created": str(workbook.properties.created) if workbook.properties.created else '',
                    "modified": str(workbook.properties.modified) if workbook.properties.modified else ''
                }
            }
            
            return result
            
        except Exception as e:
            logging.error(f"خطأ في قراءة ملف Excel: {e}")
            return {
                "file_path": file_path,
                "file_type": "excel",
                "content": "",
                "sheets": [],
                "metadata": {},
                "status": "error",
                "error": str(e)
            }
    
    def read_text_file(self, file_path: str) -> Dict[str, Any]:
        """قراءة ملف نصي"""
        try:
            result = {
                "file_path": file_path,
                "file_type": "text",
                "content": "",
                "lines": 0,
                "words": 0,
                "characters": 0,
                "metadata": {},
                "status": "success"
            }
            
            # تحديد ترميز الملف
            encodings = ['utf-8', 'utf-8-sig', 'cp1256', 'iso-8859-6', 'windows-1256']
            content = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                # محاولة قراءة كملف ثنائي
                with open(file_path, 'rb') as file:
                    content = file.read().decode('utf-8', errors='ignore')
            
            result["content"] = content
            result["lines"] = len(content.split('\n'))
            result["words"] = len(content.split())
            result["characters"] = len(content)
            
            # استخراج البيانات الوصفية
            result["metadata"] = {
                "file_size": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
            
            return result
            
        except Exception as e:
            logging.error(f"خطأ في قراءة الملف النصي: {e}")
            return {
                "file_path": file_path,
                "file_type": "text",
                "content": "",
                "lines": 0,
                "words": 0,
                "characters": 0,
                "metadata": {},
                "status": "error",
                "error": str(e)
            }
    
    def read_code_file(self, file_path: str) -> Dict[str, Any]:
        """قراءة ملف كود"""
        try:
            result = {
                "file_path": file_path,
                "file_type": "code",
                "content": "",
                "language": "",
                "lines": 0,
                "functions": 0,
                "classes": 0,
                "metadata": {},
                "status": "success"
            }
            
            # تحديد لغة البرمجة
            file_ext = Path(file_path).suffix.lower()
            language_map = {
                '.py': 'python', '.js': 'javascript', '.html': 'html', '.css': 'css',
                '.java': 'java', '.cpp': 'cpp', '.c': 'c', '.php': 'php',
                '.rb': 'ruby', '.go': 'go', '.rs': 'rust', '.swift': 'swift',
                '.kt': 'kotlin', '.scala': 'scala'
            }
            result["language"] = language_map.get(file_ext, 'unknown')
            
            # قراءة محتوى الملف
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            result["content"] = content
            result["lines"] = len(content.split('\n'))
            
            # تحليل بسيط للكود
            if result["language"] == 'python':
                result["functions"] = len(re.findall(r'def\s+\w+', content))
                result["classes"] = len(re.findall(r'class\s+\w+', content))
            elif result["language"] == 'javascript':
                result["functions"] = len(re.findall(r'function\s+\w+|const\s+\w+\s*=|let\s+\w+\s*=|var\s+\w+\s*=', content))
                result["classes"] = len(re.findall(r'class\s+\w+', content))
            elif result["language"] == 'java':
                result["functions"] = len(re.findall(r'(public|private|protected)?\s*(static)?\s*\w+\s+\w+\s*\(', content))
                result["classes"] = len(re.findall(r'class\s+\w+', content))
            
            # استخراج البيانات الوصفية
            result["metadata"] = {
                "file_size": os.path.getsize(file_path),
                "created": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                "language": result["language"]
            }
            
            return result
            
        except Exception as e:
            logging.error(f"خطأ في قراءة ملف الكود: {e}")
            return {
                "file_path": file_path,
                "file_type": "code",
                "content": "",
                "language": "",
                "lines": 0,
                "functions": 0,
                "classes": 0,
                "metadata": {},
                "status": "error",
                "error": str(e)
            }
    
    def process_file(self, file_path: str, file_size: int) -> Dict[str, Any]:
        """معالجة الملف الرئيسي"""
        try:
            # التحقق من صحة الملف
            is_valid, message = self.validate_file(file_path, file_size)
            if not is_valid:
                return {
                    "file_path": file_path,
                    "status": "error",
                    "error": message
                }
            
            # تحديد نوع الملف
            file_type = self.get_file_type(file_path)
            
            # معالجة الملف حسب نوعه
            if file_type == "pdf":
                return self.read_pdf_file(file_path)
            elif file_type == "word":
                return self.read_word_file(file_path)
            elif file_type == "excel":
                return self.read_excel_file(file_path)
            elif file_type == "text":
                return self.read_text_file(file_path)
            elif file_type == "code":
                return self.read_code_file(file_path)
            else:
                return {
                    "file_path": file_path,
                    "status": "error",
                    "error": f"نوع الملف غير مدعوم: {file_type}"
                }
                
        except Exception as e:
            logging.error(f"خطأ في معالجة الملف: {e}")
            return {
                "file_path": file_path,
                "status": "error",
                "error": str(e)
            }
    
    def save_file(self, content: str, file_name: str, file_type: str, output_dir: str = None) -> str:
        """حفظ محتوى في ملف جديد"""
        try:
            if output_dir is None:
                output_dir = self.output_folder
            
            output_path = Path(output_dir) / file_name
            
            # إضافة امتداد الملف إذا لم يكن موجوداً
            if not output_path.suffix:
                if file_type == "pdf":
                    output_path = output_path.with_suffix('.txt')  # Gemini لا ينشئ PDF مباشرة
                elif file_type == "word":
                    output_path = output_path.with_suffix('.txt')
                elif file_type == "excel":
                    output_path = output_path.with_suffix('.txt')
                elif file_type == "text":
                    output_path = output_path.with_suffix('.txt')
                elif file_type == "code":
                    output_path = output_path.with_suffix('.txt')
            
            # حفظ المحتوى
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            return str(output_path)
            
        except Exception as e:
            logging.error(f"خطأ في حفظ الملف: {e}")
            raise
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """الحصول على معلومات الملف"""
        try:
            file_stat = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            return {
                "file_name": file_path_obj.name,
                "file_size": file_stat.st_size,
                "file_type": self.get_file_type(file_path),
                "created": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                "extension": file_path_obj.suffix,
                "exists": True
            }
            
        except Exception as e:
            logging.error(f"خطأ في الحصول على معلومات الملف: {e}")
            return {
                "file_name": "",
                "file_size": 0,
                "file_type": "unknown",
                "created": "",
                "modified": "",
                "extension": "",
                "exists": False,
                "error": str(e)
            }